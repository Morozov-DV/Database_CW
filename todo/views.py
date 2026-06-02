import datetime
from django.contrib.auth import logout as auth_logout
from django.shortcuts import render, redirect
from django.views import View
from django.http import JsonResponse, HttpResponse
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.mixins import LoginRequiredMixin
from docx import Document
from .services import ProjectService, KanbanBoardService, TaskService
from .models import Employee

class DashboardView(LoginRequiredMixin, View):
    def get(self, request):
        projects = ProjectService.get_all_projects()
        return render(request, 'todo/dashboard.html', {'projects': projects})

    def post(self, request):
        project_id = request.POST.get('project_id')
        name = request.POST.get('name')
        description = request.POST.get('description')
        budget = request.POST.get('budget')

        try:
            ProjectService.create_project(
                project_id=project_id,
                name=name,
                description=description,
                budget=budget if budget else None
            )
            return redirect('dashboard')
        except Exception as e:
            projects = ProjectService.get_all_projects()
            return render(request, 'todo/dashboard.html', {
                'projects': projects,
                'error': str(e)
            })


class BoardDetailView(LoginRequiredMixin, View):
    def get(self, request, project_id):
        try:
            board_structure = KanbanBoardService.get_board_structure(project_id)
            employees = Employee.objects.all()
            return render(request, 'todo/board.html', {
                'project': board_structure['project'],
                'board_data': board_structure['board_data'],
                'employees': employees
            })
        except Exception as e:
            return redirect('dashboard')


class CreateTaskView(LoginRequiredMixin, View):
    def post(self, request, project_id):
        task_id = request.POST.get('task_id')
        name = request.POST.get('name')
        description = request.POST.get('description')
        priority = request.POST.get('priority', 'Середній')
        employee_id = request.POST.get('employee_id')
        deadline_raw = request.POST.get('deadline')

        deadline = None
        if deadline_raw:
            try:
                deadline = datetime.datetime.strptime(deadline_raw, '%Y-%m-%d')
            except ValueError:
                header_data = KanbanBoardService.get_board_structure(project_id)
                return render(request, 'todo/board.html', {
                    **header_data,
                    'employees': Employee.objects.all(),
                    'error': "Некоректний формат дати."
                })

        try:
            TaskService.create_task(
                task_id=task_id,
                project_id=project_id,
                name=name,
                description=description,
                priority=priority,
                deadline=deadline,
                employee_id=employee_id if employee_id else None
            )
            return redirect('board_detail', project_id=project_id)
        except Exception as e:
            board_structure = KanbanBoardService.get_board_structure(project_id)
            return render(request, 'todo/board.html', {
                'project': board_structure['project'],
                'board_data': board_structure['board_data'],
                'employees': Employee.objects.all(),
                'error': str(e)
            })


@method_decorator(csrf_exempt, name='dispatch')
class MoveCardView(LoginRequiredMixin, View):
    def post(self, request):
        task_id = request.POST.get('task_id')
        new_status = request.POST.get('status')

        if not task_id or not new_status:
            return JsonResponse({'success': False, 'error': 'Missing parameters'}, status=400)

        success = TaskService.move_task(task_id, new_status)
        if success:
            return JsonResponse({'success': True})
        return JsonResponse({'success': False, 'error': 'Task not found'}, status=404)


class ExportProjectReportView(LoginRequiredMixin, View):
    def get(self, request, project_id):
        try:
            board_structure = KanbanBoardService.get_board_structure(project_id)
            project = board_structure['project']
            board_data = board_structure['board_data']

            doc = Document()
            doc.add_heading(f"Звіт щодо проєкту: {project.name}", level=1)
            
            doc.add_paragraph(f"ID проєкту: {project.project_id}")
            doc.add_paragraph(f"Опис: {project.description or 'Відсутній'}")
            doc.add_paragraph(f"Бюджет: {project.budget or '0.00'} грн.")
            doc.add_paragraph(f"Фактичні витрати: {project.expenses or '0.00'} грн.")
            doc.add_paragraph(f"Дата початку: {project.start_date.strftime('%Y-%m-%d') if project.start_date else 'Не вказано'}")

            doc.add_heading("Стан виконання завдань (Канбан)", level=2)

            for status, tasks in board_data.items():
                doc.add_heading(f"Статус: {status}", level=3)
                if not tasks:
                    doc.add_paragraph("Завдання відсутні.")
                    continue
                
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'ID Завдання'
                hdr_cells[1].text = 'Назва'
                hdr_cells[2].text = 'Пріоритет'
                hdr_cells[3].text = 'Виконавець'

                for task in tasks:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(task.task_id)
                    row_cells[1].text = str(task.name)
                    row_cells[2].text = str(task.priority or 'Середній')
                    row_cells[3].text = f"{task.employee.first_name} {task.employee.last_name}" if task.employee else "Не призначено"

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=Project_{project_id}_Report.docx'
            doc.save(response)
            return response

        except Exception as e:
            return redirect('dashboard')
        
class CustomLogoutView(View):
    def get(self, request):
        auth_logout(request)
        return redirect('login')

    def post(self, request):
        auth_logout(request)
        return redirect('login')